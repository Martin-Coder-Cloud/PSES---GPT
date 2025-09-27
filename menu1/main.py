# menu1/main.py — PSES Explorer Search (Menu 1)
# UX update:
# - Hide Search/Reset after running a query; show "Start a new search" at the bottom.
# - Add "Download AI analysis (Word)" button.
# - Rename Excel button to "Download data tabulations (Excel)".
from __future__ import annotations

import io
import json
import os
import time
from datetime import datetime
from functools import lru_cache
from typing import List, Dict, Tuple, Optional

import pandas as pd
import streamlit as st

# -----------------------------
# Data loader (repo-provided)
# -----------------------------
try:
    from utils.data_loader import load_results2024_filtered  # main query function
    HAVE_LOADER = True
except Exception:
    load_results2024_filtered = None  # type: ignore
    HAVE_LOADER = False

# Optional backend info / preload hooks (best-effort)
try:
    from utils.data_loader import get_backend_info  # type: ignore
except Exception:
    def get_backend_info() -> dict:
        return {"engine": "csv.gz", "in_memory": False}

try:
    from utils.data_loader import preload_pswide_dataframe  # type: ignore
except Exception:
    def preload_pswide_dataframe():
        return None

try:
    import utils.data_loader as _dl  # type: ignore
except Exception:
    _dl = None  # type: ignore

# Hybrid search (module you created earlier); fallback if missing
try:
    from utils.hybrid_search import hybrid_question_search  # type: ignore
except Exception:
    def hybrid_question_search(qdf: pd.DataFrame, query: str, top_k: int = 120, min_score: float = 0.40) -> pd.DataFrame:
        if not query or not str(query).strip():
            return pd.DataFrame(columns=["code", "text", "display", "score"])
        q = str(query).strip().lower()
        tokens = {t for t in q.replace(",", " ").split() if t}
        scores = []
        for _, r in qdf.iterrows():
            text = f"{r['code']} {r['text']}".lower()
            base = 1.0 if q in text else 0.0
            overlap = sum(1 for t in tokens if t in text) / max(len(tokens), 1)
            score = 0.6 * overlap + 0.4 * base
            scores.append(score)
        out = qdf.copy()
        out["score"] = scores
        out = out.sort_values("score", ascending=False)
        out = out[out["score"] >= min_score]
        return out.head(top_k)

# -----------------------------
# OpenAI (best-effort; wrapped)
# -----------------------------
OPENAI_API_KEY = st.secrets.get("OPENAI_API_KEY", "") or os.environ.get("OPENAI_API_KEY", "")
OPENAI_MODEL = st.secrets.get("OPENAI_MODEL", "gpt-4o-mini")

def _call_openai_json(system: str, user: str, model: str = OPENAI_MODEL, temperature: float = 0.2, max_retries: int = 2):
    """Return (json_text, error_hint). Also logs health in session_state."""
    if not OPENAI_API_KEY:
        st.session_state["menu1_last_ai_status"] = {"time": datetime.now().isoformat(timespec="seconds"),
                                                    "ok": False, "hint": "no_api_key"}
        return "", "no_api_key"
    try:
        from openai import OpenAI  # type: ignore
        client = OpenAI(api_key=OPENAI_API_KEY)
        hint = "unknown_error"
        for attempt in range(max_retries + 1):
            try:
                resp = client.chat.completions.create(
                    model=model,
                    temperature=temperature,
                    response_format={"type": "json_object"},
                    messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                )
                content = resp.choices[0].message.content or ""
                st.session_state["menu1_last_ai_status"] = {"time": datetime.now().isoformat(timespec="seconds"),
                                                            "ok": True, "hint": None}
                return content, None
            except Exception as e:
                hint = f"openai_err_{attempt+1}: {type(e).__name__}"
                time.sleep(0.8 * (attempt + 1))
        st.session_state["menu1_last_ai_status"] = {"time": datetime.now().isoformat(timespec="seconds"),
                                                    "ok": False, "hint": hint}
        return "", hint
    except Exception:
        try:
            import openai  # type: ignore
            openai.api_key = OPENAI_API_KEY
            hint = "unknown_error"
            for attempt in range(max_retries + 1):
                try:
                    resp = openai.ChatCompletion.create(
                        model=model,
                        temperature=temperature,
                        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                    )
                    content = resp["choices"][0]["message"]["content"] or ""
                    st.session_state["menu1_last_ai_status"] = {"time": datetime.now().isoformat(timespec="seconds"),
                                                                "ok": True, "hint": None}
                    return content, None
                except Exception as e:
                    hint = f"openai_legacy_err_{attempt+1}: {type(e).__name__}"
                    time.sleep(0.8 * (attempt + 1))
            st.session_state["menu1_last_ai_status"] = {"time": datetime.now().isoformat(timespec="seconds"),
                                                        "ok": False, "hint": hint}
            return "", hint
        except Exception:
            st.session_state["menu1_last_ai_status"] = {"time": datetime.now().isoformat(timespec="seconds"),
                                                        "ok": False, "hint": "no_openai_sdk"}
            return "", "no_openai_sdk"

# -----------------------------
# Exact AI system prompt (unchanged)
# -----------------------------
AI_SYSTEM_PROMPT = (
    "You are preparing insights for the Government of Canada’s Public Service Employee Survey (PSES).\n\n"
    "Context\n"
    "- The PSES provides information to improve people management practices in the federal public service.\n"
    "- Results help departments and agencies identify strengths and concerns in areas such as employee engagement, anti-racism, equity and inclusion, and workplace well-being.\n"
    "- The survey tracks progress over time to refine action plans. Employees’ voices guide improvements to workplace quality, which leads to better results for the public service and Canadians.\n"
    "- Each cycle includes recurring questions (for tracking trends) and new/modified questions reflecting evolving priorities (e.g., updated Employment Equity questions and streamlined hybrid-work items in 2024).\n"
    "- Statistics Canada administers the survey with the Treasury Board of Canada Secretariat. Confidentiality is guaranteed under the Statistics Act (grouped reporting; results for groups <10 are suppressed).\n\n"
    "Data-use rules (hard constraints)\n"
    "- Use ONLY the provided JSON payload/table. DO NOT invent, assume, extrapolate, infer, or generalize beyond the numbers present. No speculation or hypotheses.\n"
    "- Public Service–wide scope ONLY; do not reference specific departments unless present in the payload.\n"
    "- Express percentages as whole numbers (e.g., “75%”). Use “points” for differences/changes.\n\n"
    "Analysis rules\n"
    "- Begin with the 2024 result for the selected question (metric_label).\n"
    "- Describe trend over time: compare 2024 with the earliest year available, using thresholds:\n"
    "  • stable ≤1 point\n"
    "  • slight >1–2 points\n"
    "  • notable >2 points\n"
    "- Compare demographic groups in 2024:\n"
    "  • Focus on the most relevant comparisons (largest gap(s), or those crossing thresholds).\n"
    "  • Report gaps in points and classify them: minimal ≤2, notable >2–5, important >5.\n"
    "- If multiple groups are present, highlight only the most meaningful contrasts instead of exhaustively listing all.\n"
    "- Mention whether gaps observed in 2024 have widened, narrowed, or remained stable compared with earlier years.\n"
    "- Conclude with a concise overall statement (e.g., “Overall, results have remained steady and demographic gaps are unchanged”).\n\n"
    "Style & output\n"
    "- Professional, concise, neutral. Narrative style (1–3 short paragraphs, no lists).\n"
    "- Output VALID JSON with exactly one key: \"narrative\".\n"
)

# -----------------------------
# Cached metadata
# -----------------------------
@st.cache_data(show_spinner=False)
def _load_demographics() -> pd.DataFrame:
    df = pd.read_excel("metadata/Demographics.xlsx")
    df.columns = [c.strip() for c in df.columns]
    return df

@st.cache_data(show_spinner=False)
def _load_questions() -> pd.DataFrame:
    qdf = pd.read_excel("metadata/Survey Questions.xlsx")
    qdf.columns = [c.strip().lower() for c in qdf.columns]
    if "question" in qdf.columns and "english" in qdf.columns:
        qdf = qdf.rename(columns={"question": "code", "english": "text"})
    qdf["code"] = qdf["code"].astype(str)
    qdf["qnum"] = qdf["code"].str.extract(r"Q?(\d+)", expand=False)
    with pd.option_context("mode.chained_assignment", None):
        qdf["qnum"] = pd.to_numeric(qdf["qnum"], errors="coerce")
    qdf = qdf.sort_values(["qnum", "code"], na_position="last")
    qdf["display"] = qdf["code"].astype(str) + " – " + qdf["text"].astype(str)
    return qdf[["code", "text", "display"]]

@st.cache_data(show_spinner=False)
def _load_scales() -> pd.DataFrame:
    sdf = pd.read_excel("metadata/Survey Scales.xlsx")
    sdf.columns = [c.strip().lower() for c in sdf.columns]
    return sdf

# -----------------------------
# Helpers (demographics / display / summary)
# -----------------------------
def _is_overall(val) -> bool:
    if val is None:
        return True
    s = str(val).strip().lower()
    return s in ("", "all", "all respondents", "allrespondents")

def _resolve_demcodes(demo_df: pd.DataFrame, category_label: str, subgroup_label: Optional[str]):
    DEMO_CAT_COL = "DEMCODE Category"
    LABEL_COL = "DESCRIP_E"
    if not category_label or category_label == "All respondents":
        return [None], {None: "All respondents"}, False

    code_col = None
    for c in ["DEMCODE", "DemCode", "CODE", "Code", "CODE_E", "Demographic code"]:
        if c in demo_df.columns:
            code_col = c
            break

    df_cat = demo_df[demo_df[DEMO_CAT_COL] == category_label] if DEMO_CAT_COL in demo_df.columns else demo_df.copy()
    if df_cat.empty:
        return [None], {None: "All respondents"}, False

    if subgroup_label:
        if code_col and LABEL_COL in df_cat.columns:
            r = df_cat[df_cat[LABEL_COL] == subgroup_label]
            if not r.empty:
                code = str(r.iloc[0][code_col])
                return [code], {code: subgroup_label}, True
        return [subgroup_label], {subgroup_label: subgroup_label}, True

    if code_col and LABEL_COL in df_cat.columns:
        codes = df_cat[code_col].astype(str).tolist()
        labels = df_cat[LABEL_COL].astype(str).tolist()
        keep = [(c, l) for c, l in zip(codes, labels) if str(c).strip() != ""]
        codes = [c for c, _ in keep]
        disp_map = {c: l for c, l in keep}
        return codes, disp_map, True

    if LABEL_COL in df_cat.columns:
        labels = df_cat[LABEL_COL].astype(str).tolist()
        return labels, {l: l for l in labels}, True

    return [None], {None: "All respondents"}, False

@lru_cache(maxsize=512)
def _get_scale_labels_cached(qcode_upper: str) -> Tuple[Tuple[str, str], ...]:
    sdf = _load_scales()
    candidates = pd.DataFrame()
    for key in ["code", "question"]:
        if key in sdf.columns:
            candidates = sdf[sdf[key].astype(str).str.upper() == qcode_upper]
            if not candidates.empty:
                break
    pairs: List[Tuple[str, str]] = []
    for i in range(1, 8):
        col = f"answer{i}"
        lbl = None
        if not candidates.empty and col in sdf.columns:
            vals = candidates[col].dropna().astype(str)
            if not vals.empty:
                lbl = vals.iloc[0].strip()
        pairs.append((col, lbl or f"Answer {i}"))
    return tuple(pairs)

def _get_scale_labels(scales_df: pd.DataFrame, question_code: str):
    return list(_get_scale_labels_cached(str(question_code).upper()))

def _drop_999(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    out = df.copy()
    for c in [f"answer{i}" for i in range(1, 8)] + ["POSITIVE","NEUTRAL","NEGATIVE","AGREE","ANSCOUNT","positive_pct","neutral_pct","negative_pct","n"]:
        if c in out.columns:
            v = pd.to_numeric(out[c], errors="coerce")
            out.loc[v.isin([999, 9999]), c] = pd.NA
    return out

def _normalize_results(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    if "question_code" not in out.columns:
        if "QUESTION" in out.columns: out = out.rename(columns={"QUESTION": "question_code"})
        else:
            for c in out.columns:
                if c.strip().lower() == "question":
                    out = out.rename(columns={c: "question_code"}); break
    if "year" not in out.columns:
        if "SURVEYR" in out.columns: out = out.rename(columns={"SURVEYR": "year"})
        else:
            for c in out.columns:
                if c.strip().lower() in ("surveyr","year"):
                    out = out.rename(columns={c: "year"}); break
    if "group_value" not in out.columns:
        if "DEMCODE" in out.columns: out = out.rename(columns={"DEMCODE": "group_value"})
        else:
            for c in out.columns:
                if c.strip().lower() == "demcode":
                    out = out.rename(columns={c: "group_value"}); break
    if "positive_pct" not in out.columns and "POSITIVE" in out.columns:
        out = out.rename(columns={"POSITIVE": "positive_pct"})
    if "neutral_pct" not in out.columns and "NEUTRAL" in out.columns:
        out = out.rename(columns={"NEUTRAL": "neutral_pct"})
    if "negative_pct" not in out.columns and "NEGATIVE" in out.columns:
        out = out.rename(columns={"NEGATIVE": "negative_pct"})
    if "AGREE" in out.columns and "Agree" not in out.columns:
        out = out.rename(columns={"AGREE": "Agree"})
    if "n" not in out.columns and "ANSCOUNT" in out.columns:
        out = out.rename(columns={"ANSCOUNT": "n"})
    return out

def _format_display(df_slice: pd.DataFrame, dem_disp_map: Dict, category_in_play: bool, scale_pairs: List[Tuple[str,str]]) -> pd.DataFrame:
    if df_slice.empty:
        return df_slice.copy()
    out = df_slice.copy()
    out["YearNum"] = pd.to_numeric(out["year"], errors="coerce").astype("Int64")
    out["Year"] = out["YearNum"].astype(str)

    if category_in_play:
        def to_label(code):
            if _is_overall(code):
                return "All respondents"
            return dem_disp_map.get(code, dem_disp_map.get(str(code), str(code)))
        out["Demographic"] = out["group_value"].apply(to_label)

    dist_cols = [k for k,_ in scale_pairs if k in out.columns]
    rename_map = {k: v for k, v in scale_pairs if k in out.columns}

    keep_cols = ["YearNum","Year"] + (["Demographic"] if category_in_play else []) + dist_cols + ["positive_pct","neutral_pct","negative_pct","Agree","n"]
    keep_cols = [c for c in keep_cols if c in out.columns]
    out = out[keep_cols].rename(columns=rename_map).copy()
    out = out.rename(columns={"positive_pct":"Positive","neutral_pct":"Neutral","negative_pct":"Negative"})

    sort_cols = ["YearNum"] + (["Demographic"] if category_in_play else [])
    sort_asc = [False] + ([True] if category_in_play else [])
    out = out.sort_values(sort_cols, ascending=sort_asc, kind="mergesort").reset_index(drop=True)
    out = out.drop(columns=["YearNum"])

    for c in out.columns:
        if c not in ("Year","Demographic"):
            out[c] = pd.to_numeric(out[c], errors="coerce")
    pct_like = [c for c in out.columns if c not in ("Year","Demographic","n")]
    if pct_like:
        out[pct_like] = out[pct_like].round(1)
    if "n" in out.columns:
        out["n"] = pd.to_numeric(out["n"], errors="coerce").astype("Int64")
    return out

def _detect_metric_presence(df_disp: pd.DataFrame, scale_pairs: List[Tuple[str,str]]) -> Dict[str, Tuple[bool, Optional[str]]]:
    res: Dict[str, Tuple[bool, Optional[str]]] = {
        "Positive": (False, "% positive"),
        "Agree": (False, "% agree"),
        "Answer1": (False, None),
    }
    cols_low = {c.lower(): c for c in df_disp.columns}
    if "positive" in cols_low:
        has = pd.to_numeric(df_disp[cols_low["positive"]], errors="coerce").notna().any()
        res["Positive"] = (has, "% positive")
    if "agree" in cols_low:
        has = pd.to_numeric(df_disp[cols_low["agree"]], errors="coerce").notna().any()
        res["Agree"] = (has, "% agree")
    ans1_label = None
    for k, lbl in scale_pairs:
        ans1_label = lbl or "Answer 1"
        if lbl and lbl in df_disp.columns:
            has = pd.to_numeric(df_disp[lbl], errors="coerce").notna().any()
            res["Answer1"] = (bool(has), f"% {lbl}")
            break
        else:
            if "answer1" in cols_low:
                has = pd.to_numeric(df_disp[cols_low["answer1"]], errors="coerce").notna().any()
                res["Answer1"] = (bool(has), f"% {ans1_label}")
                break
    if res["Answer1"][1] is None and ans1_label:
        res["Answer1"] = (res["Answer1"][0], f"% {ans1_label}")
    return res

# -----------------------------
# State reset
# -----------------------------
def _delete_keys(prefixes: List[str], exact_keys: List[str] = None):
    exact_keys = exact_keys or []
    for k in list(st.session_state.keys()):
        if any(k.startswith(p) for p in prefixes) or (k in exact_keys):
            try: del st.session_state[k]
            except Exception: pass

def _reset_menu1_state():
    year_keys = [f"year_{y}" for y in (2024, 2022, 2020, 2019)]
    exact = [
        "menu1_selected_codes","menu1_selected_order","menu1_hits","menu1_kw_query","menu1_last_kw",
        "menu1_multi_questions","menu1_ai_toggle","menu1_show_diag","select_all_years","demo_main",
        "menu1_find_hits","last_query_info","menu1_last_ai_status","menu1_search_clicked"
    ] + year_keys
    prefixes = ["kwhit_","sel_","sub_"]
    _delete_keys(prefixes, exact)
    st.session_state.setdefault("menu1_kw_query", "")
    st.session_state.setdefault("menu1_last_kw", None)
    st.session_state.setdefault("menu1_hits", [])
    st.session_state.setdefault("menu1_selected_codes", [])
    st.session_state.setdefault("menu1_selected_order", [])
    st.session_state.setdefault("menu1_multi_questions", [])
    st.session_state.setdefault("menu1_ai_toggle", True)
    st.session_state.setdefault("menu1_show_diag", False)
    st.session_state.setdefault("last_query_info", None)
    st.session_state.setdefault("menu1_last_ai_status", None)
    st.session_state.setdefault("menu1_search_clicked", False)

# -----------------------------
# AI payload builders
# -----------------------------
def _series_json(df_disp: pd.DataFrame, metric_col: str) -> List[Dict[str, float]]:
    rows = []
    s = df_disp.copy()
    if "Demographic" in s.columns:
        s = s.groupby("Year", as_index=False)[metric_col].mean(numeric_only=True)
        s = s.rename(columns={metric_col: "Metric"})
    else:
        s = s[["Year", metric_col]].rename(columns={metric_col: "Metric"})
    s = s.dropna(subset=["Year"]).sort_values("Year")
    for _, r in s.iterrows():
        try: y = int(r["Year"])
        except Exception: y = r["Year"]
        rows.append({"year": y, "value": float(r["Metric"]) if pd.notna(r["Metric"]) else None})
    return rows

def _user_prompt_per_q(qcode: str, qtext: str, df_disp: pd.DataFrame, metric_col: str, metric_label: str, category_in_play: bool) -> str:
    latest = pd.to_numeric(df_disp["Year"], errors="coerce").max()
    group_info = []
    if category_in_play and "Demographic" in df_disp.columns and pd.notna(latest):
        g = df_disp[pd.to_numeric(df_disp["Year"], errors="coerce") == latest][["Demographic", metric_col]].dropna()
        g = g.sort_values(metric_col, ascending=False)
        if not g.empty:
            top = g.iloc[0].to_dict(); bot = g.iloc[-1].to_dict()
            group_info = [
                {"demographic": str(top["Demographic"]), "value": float(top[metric_col])},
                {"demographic": str(bot["Demographic"]), "value": float(bot[metric_col])},
            ]
    payload = {
        "question_code": qcode,
        "question_text": qtext,
        "metric_label": metric_label,
        "series_positive_by_year": _series_json(df_disp, metric_col),
        "latest_year_group_snapshot": group_info,
        "notes": "Use the supplied metric_label for interpretation."
    }
    return json.dumps(payload, ensure_ascii=False)

def _user_prompt_overall_mixed(per_q_disp: Dict[str, pd.DataFrame],
                               per_q_text: Dict[str, str],
                               per_q_metric_col: Dict[str, str],
                               per_q_metric_label: Dict[str, str]) -> str:
    items = []
    for qcode, df in per_q_disp.items():
        metric_col = per_q_metric_col[qcode]
        metric_label = per_q_metric_label[qcode]
        s = df.copy()
        if "Demographic" in s.columns:
            s = s.groupby("Year", as_index=False)[metric_col].mean(numeric_only=True)
            s = s.rename(columns={metric_col: "Metric"})
        else:
            s = s[["Year", metric_col]].rename(columns={metric_col: "Metric"})
        s["Year"] = pd.to_numeric(s["Year"], errors="coerce")
        series = {}
        for _, r in s.dropna(subset=["Year"]).sort_values("Year").iterrows():
            series[int(r["Year"])] = float(r["Metric"]) if pd.notna(r["Metric"]) else None
        items.append({
            "question_label": f"{qcode} — {per_q_text.get(qcode,'')}".strip().rstrip(" —"),
            "metric_label": metric_label,
            "values_by_year": series
        })
    return json.dumps({"questions": items, "notes": "Provide a concise overall synthesis across questions, respecting each question’s metric_label."}, ensure_ascii=False)

# -----------------------------
# Misc utils
# -----------------------------
def _get_pswide_df() -> Optional[pd.DataFrame]:
    try:
        df = preload_pswide_dataframe()
        if isinstance(df, pd.DataFrame):
            return df
    except Exception:
        pass
    try:
        if _dl and hasattr(_dl, "pswide_df"):
            df = getattr(_dl, "pswide_df")
            if isinstance(df, pd.DataFrame):
                return df
    except Exception:
        pass
    return None

def _kv_table(d: Dict[str, object]) -> pd.DataFrame:
    rows = [(k, (", ".join(map(str, v)) if isinstance(v, (list, tuple)) else v)) for k, v in d.items()]
    return pd.DataFrame(rows, columns=["Field", "Value"])

# -----------------------------
# UI
# -----------------------------
def run_menu1():
    st.markdown("""
        <style>
            body { background-image: none !important; background-color: white !important; }
            .block-container { padding-top: 1rem !important; }
            .menu-banner { width: 100%; height: auto; display: block; margin-top: 0px; margin-bottom: 20px; }
            .custom-header { font-size: 30px !important; font-weight: 700; margin-bottom: 6px; }
            .custom-instruction { font-size: 16px !important; line-height: 1.4; margin-bottom: 10px; color: #333; }
            .field-label { font-size: 18px !important; font-weight: 600 !important; margin-top: 12px !important; margin-bottom: 2px !important; color: #222 !important; }
            .action-row { display:flex; gap:10px; align-items:center; }
            [data-testid="stSwitch"] div[role="switch"][aria-checked="true"] { background-color: #e03131 !important; }
            [data-testid="stSwitch"] div[role="switch"] { box-shadow: inset 0 0 0 1px rgba(0,0,0,0.1); }
            .tiny-note { font-size: 13px; color: #444; margin-bottom: 6px; }
            .diag-box { background: #fafafa; border: 1px solid #eee; border-radius: 8px; padding: 10px 12px; }
            /* Ensure tabs and dataframes fill the center column width */
            div[data-baseweb="tab-panel"] { width: 100% !important; }
            div[data-testid="stDataFrame"] { width: 100% !important; }
        </style>
    """, unsafe_allow_html=True)

    if st.session_state.get("last_active_menu") != "menu1":
        _reset_menu1_state()
    st.session_state["last_active_menu"] = "menu1"

    # Defaults BEFORE widget creation
    st.session_state.setdefault("menu1_selected_codes", [])
    st.session_state.setdefault("menu1_selected_order", [])
    st.session_state.setdefault("menu1_hits", [])
    st.session_state.setdefault("menu1_kw_query", "")
    st.session_state.setdefault("menu1_last_kw", None)
    st.session_state.setdefault("menu1_multi_questions", [])
    st.session_state.setdefault("menu1_ai_toggle", True)
    st.session_state.setdefault("menu1_show_diag", False)
    st.session_state.setdefault("last_query_info", None)
    st.session_state.setdefault("menu1_last_ai_status", None)
    st.session_state.setdefault("menu1_search_clicked", False)

    # Early guard
    if load_results2024_filtered is None:
        st.error("Data loader unavailable. Please verify utils.data_loader.load_results2024_filtered.")
    loader_ok = load_results2024_filtered is not None

    demo_df = _load_demographics()
    qdf = _load_questions()
    sdf = _load_scales()

    code_to_text = dict(zip(qdf["code"], qdf["text"]))
    code_to_display = dict(zip(qdf["code"], qdf["display"]))
    display_to_code = {v: k for k, v in code_to_display.items()}

    left, center, right = st.columns([1, 3, 1])
    with center:
        # Banner
        st.markdown(
            "<img class='menu-banner' src='https://raw.githubusercontent.com/Martin-Coder-Cloud/PSES---GPT/refs/heads/main/PSES%20email%20banner.png'>",
            unsafe_allow_html=True
        )

        # Title
        st.markdown('<div class="custom-header">PSES Explorer Search</div>', unsafe_allow_html=True)

        # Toggles row
        c1, c2 = st.columns([1, 1])
        with c1:
            st.toggle("🧠 Enable AI analysis", key="menu1_ai_toggle", help="Include the AI-generated analysis alongside the tables.")
        with c2:
            st.toggle("🔧 Show technical parameters & diagnostics", key="menu1_show_diag", help="Show current parameters, app setup status, and last query timings.")

        # Instructions
        st.markdown("""
            <div class="custom-instruction">
                Please use this menu to explore the survey results by questions.<br>
                You may select from the drop down menu below up to five questions or find questions via the keyword/theme search.
                Select year(s) and optionally a demographic category and subgroup.
            </div>
        """, unsafe_allow_html=True)

        # Diagnostics panel (tabbed tables)
        if st.session_state.get("menu1_show_diag", False):
            tabs_diag = st.tabs(["Loading status", "Parameters", "AI health check", "Last query"])
            with tabs_diag[0]:
                info = {}
                try: info = get_backend_info() or {}
                except Exception: info = {"engine": "csv.gz"}
                try:
                    df_ps = _get_pswide_df()
                    if isinstance(df_ps, pd.DataFrame) and not df_ps.empty:
                        ycol = "year" if "year" in df_ps.columns else ("SURVEYR" if "SURVEYR" in df_ps.columns else None)
                        qcol = "question_code" if "question_code" in df_ps.columns else ("QUESTION" if "QUESTION" in df_ps.columns else None)
                        yr_min = int(pd.to_numeric(df_ps[ycol], errors="coerce").min()) if ycol else None
                        yr_max = int(pd.to_numeric(df_ps[ycol], errors="coerce").max()) if ycol else None
                        info.update({"in_memory": True, "pswide_rows": int(len(df_ps)),
                                     "unique_questions": int(df_ps[qcol].astype(str).nunique()) if qcol else None,
                                     "year_range": f"{yr_min}–{yr_max}" if yr_min and yr_max else None})
                    else:
                        info.update({"in_memory": False})
                except Exception:
                    pass
                try: info["metadata_questions"] = int(len(qdf))
                except Exception: pass
                try: info["metadata_scales"] = int(len(sdf))
                except Exception: pass
                try: info["metadata_demographics"] = int(len(demo_df))
                except Exception: pass
                st.table(_kv_table(info))
            with tabs_diag[1]:
                years_selected = [y for y in [2019, 2020, 2022, 2024] if st.session_state.get(f"year_{y}", True if st.session_state.get("select_all_years", True) else False)]
                demo_cat = st.session_state.get("demo_main", "All respondents")
                subkey = f"sub_{str(demo_cat).replace(' ', '_')}"
                subgroup = ("All respondents" if not demo_cat or demo_cat == "All respondents"
                            else (st.session_state.get(subkey, "") or "(all subgroups)"))
                params = {"Selected questions": [code_to_display.get(c, c) for c in st.session_state.get("menu1_selected_codes", [])],
                          "Years selected": years_selected, "Demographic category": demo_cat or "All respondents", "Subgroup": subgroup}
                st.table(_kv_table(params))
            with tabs_diag[2]:
                ai_status = st.session_state.get("menu1_last_ai_status") or {}
                health = {"OpenAI model": OPENAI_MODEL, "API key configured": bool(OPENAI_API_KEY),
                          "Last AI run time": ai_status.get("time"), "Last AI ok": ai_status.get("ok"),
                          "Last AI hint": ai_status.get("hint")}
                st.table(_kv_table(health))
            with tabs_diag[3]:
                last = st.session_state.get("last_query_info") or {"status": "No query yet"}
                st.table(_kv_table(last))

        # ---------- Question selection ----------
        st.markdown('<div class="field-label">Pick up to 5 survey questions:</div>', unsafe_allow_html=True)
        all_displays = qdf["display"].tolist()
        multi_choices = st.multiselect(
            "Choose one or more from the official list",
            all_displays,
            default=st.session_state.get("menu1_multi_questions", []),
            max_selections=5,
            label_visibility="collapsed",
            key="menu1_multi_questions",
        )
        selected_from_multi: List[str] = [display_to_code[d] for d in multi_choices if d in display_to_code]

        with st.expander("Search by keywords or theme (optional)"):
            search_query = st.text_input("Enter keywords (e.g., harassment, recognition, onboarding)", key="menu1_kw_query")
            if st.button("Search questions", key="menu1_find_hits"):
                if st.session_state.get("menu1_last_kw") == (search_query or "").strip():
                    st.info("Same keyword as last search; skipped re-running.")
                else:
                    hits_df = hybrid_question_search(qdf, search_query, top_k=120, min_score=0.40)
                    st.session_state["menu1_hits"] = hits_df[["code", "text"]].to_dict(orient="records") if not hits_df.empty else []
                    st.session_state["menu1_last_kw"] = (search_query or "").strip()

            selected_from_hits: List[str] = []
            if st.session_state["menu1_hits"]:
                st.write(f"Top {len(st.session_state['menu1_hits'])} matches meeting the quality threshold:")
                for rec in st.session_state["menu1_hits"]:
                    code = rec["code"]; text = rec["text"]
                    label = f"{code} — {text}"
                    key = f"kwhit_{code}"
                    default_checked = st.session_state.get(key, False) or (code in selected_from_multi)
                    checked = st.checkbox(label, value=default_checked, key=key)
                    if checked:
                        selected_from_hits.append(code)
            else:
                st.info('Enter keywords and click "Search questions" to see matches.')

        # Stable order merge
        current_selected = selected_from_multi + [c for c in selected_from_hits if c not in selected_from_multi]
        prev_order: List[str] = st.session_state.get("menu1_selected_order", [])
        new_order = [c for c in prev_order if c in current_selected]
        for c in current_selected:
            if c not in new_order:
                new_order.append(c)
        if len(new_order) > 5:
            new_order = new_order[:5]
            st.warning("Limit is 5 questions; extra selections were ignored.")
        st.session_state["menu1_selected_order"] = new_order
        st.session_state["menu1_selected_codes"] = new_order

        if st.session_state["menu1_selected_codes"]:
            st.markdown('<div class="field-label">Selected questions:</div>', unsafe_allow_html=True)
            updated = list(st.session_state["menu1_selected_codes"])
            cols = st.columns(min(5, len(updated)))
            for idx, code in enumerate(list(updated)):
                with cols[idx % len(cols)]:
                    label = code_to_display.get(code, code)
                    keep = st.checkbox(label, value=True, key=f"sel_{code}")
                    if not keep:
                        updated = [c for c in updated if c != code]
                        hk = f"kwhit_{code}"
                        if hk in st.session_state: st.session_state[hk] = False
                        disp = code_to_display.get(code)
                        if disp:
                            st.session_state["menu1_multi_questions"] = [d for d in st.session_state["menu1_multi_questions"] if d != disp]
            if updated != st.session_state["menu1_selected_codes"]:
                st.session_state["menu1_selected_codes"] = updated
                st.session_state["menu1_selected_order"] = updated

        question_codes: List[str] = st.session_state["menu1_selected_codes"]

        # ---------- Years ----------
        st.markdown('<div class="field-label">Select survey year(s):</div>', unsafe_allow_html=True)
        all_years = [2024, 2022, 2020, 2019]
        st.session_state.setdefault("select_all_years", True)
        select_all = st.checkbox("All years", key="select_all_years")
        selected_years: List[int] = []
        year_cols = st.columns(len(all_years))
        for idx, yr in enumerate(all_years):
            with year_cols[idx]:
                default_checked = True if select_all else st.session_state.get(f"year_{yr}", False)
                if st.checkbox(str(yr), value=default_checked, key=f"year_{yr}"):
                    selected_years.append(yr)
        selected_years = sorted(selected_years)
        no_years = (len(selected_years) == 0)
        if not select_all and no_years:
            st.caption("Hint: choose at least one year, or turn on “All years”.")

        # ---------- Demographics ----------
        st.markdown('<div class="field-label">Select a demographic category (optional):</div>', unsafe_allow_html=True)
        DEMO_CAT_COL = "DEMCODE Category"
        LABEL_COL = "DESCRIP_E"
        demo_categories = ["All respondents"] + sorted(demo_df[DEMO_CAT_COL].dropna().astype(str).unique().tolist())
        st.session_state.setdefault("demo_main", "All respondents")
        demo_selection = st.selectbox("Demographic category", demo_categories, key="demo_main", label_visibility="collapsed")

        sub_selection = None
        if demo_selection != "All respondents":
            st.markdown(f'<div class="field-label">Subgroup ({demo_selection}) (optional):</div>', unsafe_allow_html=True)
            sub_items = demo_df.loc[demo_df[DEMO_CAT_COL] == demo_selection, LABEL_COL].dropna().astype(str).unique().tolist()
            sub_items = sorted(sub_items)
            sub_key = f"sub_{demo_selection.replace(' ', '_')}"
            sub_selection = st.selectbox("(leave blank to include all subgroups in this category)", [""] + sub_items, key=sub_key, label_visibility="collapsed")
            if sub_selection == "":
                sub_selection = None

        # ---------- Action row (hidden after search) ----------
        if not st.session_state.get("menu1_search_clicked", False):
            st.markdown("<div class='action-row'>", unsafe_allow_html=True)
            colA, colB = st.columns([1.2, 1])
            with colA:
                reasons = []
                if not loader_ok: reasons.append("data loader unavailable")
                if not question_codes: reasons.append("pick at least 1 question")
                if no_years: reasons.append("pick at least 1 year")
                disable_search = bool(reasons)
                label = "Search" if not disable_search else f"Search (disabled: {', '.join(reasons)})"
                clicked = st.button(label, disabled=disable_search, key="menu1_search_btn")
            with colB:
                if st.button("Reset all parameters"):
                    _reset_menu1_state()
                    st.experimental_rerun()
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            clicked = False  # already searched; keep rendering results below

        # =========================
        # RESULTS & AI (FULL WIDTH)
        # =========================
        if clicked or st.session_state.get("menu1_search_clicked", False):
            if clicked:
                st.session_state["menu1_search_clicked"] = True  # hide action row on subsequent renders

            t0 = time.time()
            with st.spinner(f"Running query… {datetime.fromtimestamp(t0).strftime('%Y-%m-%d %H:%M:%S')}"):
                demcodes, disp_map, category_in_play = _resolve_demcodes(demo_df, demo_selection, sub_selection)

                per_q_disp: Dict[str, pd.DataFrame] = {}
                per_q_text: Dict[str, str] = {}
                per_q_metric_presence: Dict[str, Dict[str, Tuple[bool, Optional[str]]]] = {}
                per_q_ans1_label: Dict[str, Optional[str]] = {}
                per_q_best_col: Dict[str, str] = {}
                per_q_best_label: Dict[str, str] = {}

                for qcode in question_codes:
                    qtext = code_to_text.get(qcode, "")
                    per_q_text[qcode] = qtext

                    parts = []
                    if load_results2024_filtered is None:
                        continue
                    for code in demcodes:
                        df_part = load_results2024_filtered(
                            question_code=qcode,
                            years=selected_years,
                            group_value=(None if _is_overall(code) else str(code))
                        )
                        if df_part is not None and not df_part.empty:
                            parts.append(df_part)
                    if not parts:
                        continue
                    df_all = pd.concat(parts, ignore_index=True)

                    df_all = _normalize_results(df_all)
                    qmask = df_all["question_code"].astype(str).str.strip().str.upper() == str(qcode).strip().upper()
                    ymask = pd.to_numeric(df_all["year"], errors="coerce").astype("Int64").isin(selected_years)
                    if demo_selection == "All respondents":
                        gv = df_all["group_value"].astype(str).fillna("").str.strip()
                        gmask = gv.apply(_is_overall)
                    else:
                        gmask = df_all["group_value"].astype(str).isin([str(c) for c in demcodes])
                    df_all = df_all[qmask & ymask & gmask].copy()
                    df_all = _drop_999(df_all)
                    if df_all.empty:
                        continue

                    scale_pairs = _get_scale_labels(sdf, qcode)
                    df_disp = _format_display(df_slice=df_all, dem_disp_map=disp_map, category_in_play=category_in_play, scale_pairs=scale_pairs)
                    if df_disp.empty:
                        continue

                    presence = _detect_metric_presence(df_disp, scale_pairs)
                    per_q_metric_presence[qcode] = presence
                    per_q_ans1_label[qcode] = presence["Answer1"][1]

                    if presence["Positive"][0]:
                        per_q_best_col[qcode] = "Positive"; per_q_best_label[qcode] = "% positive"
                    elif presence["Agree"][0]:
                        per_q_best_col[qcode] = "Agree"; per_q_best_label[qcode] = "% agree"
                    elif presence["Answer1"][0]:
                        lbl = (per_q_ans1_label[qcode] or "Answer 1")
                        per_q_best_col[qcode] = lbl if lbl in df_disp.columns else ("Answer 1" if "Answer 1" in df_disp.columns else lbl)
                        per_q_best_label[qcode] = f"% {per_q_ans1_label[qcode] or 'Answer 1'}"
                    else:
                        per_q_best_col[qcode] = "Positive"; per_q_best_label[qcode] = "% positive"

                    per_q_disp[qcode] = df_disp

            t1 = time.time()
            st.session_state["last_query_info"] = {
                "started": datetime.fromtimestamp(t0).strftime("%Y-%m-%d %H:%M:%S"),
                "finished": datetime.fromtimestamp(t1).strftime("%Y-%m-%d %H:%M:%S"),
                "elapsed_seconds": round(t1 - t0, 2),
                "engine": (get_backend_info() or {}).get("engine", "unknown"),
                "questions": len([qc for qc in question_codes if qc in per_q_disp]),
                "years": selected_years,
            }

            if not per_q_disp:
                st.info("No data found for your selection.")
            else:
                # Common summary metric decision
                qlist = [qc for qc in question_codes if qc in per_q_disp]
                def all_have(kind: str) -> bool:
                    for qc in qlist:
                        has, _ = per_q_metric_presence[qc][kind]
                        if not has: return False
                    return True
                summary_metric_col = None
                summary_metric_label = None
                if all_have("Positive"):
                    summary_metric_col = "Positive"; summary_metric_label = "% Positive"
                elif all_have("Agree"):
                    summary_metric_col = "Agree"; summary_metric_label = "% Agree"
                elif all_have("Answer1"):
                    summary_metric_col = "Answer1Value"; summary_metric_label = "% Answer 1"
                else:
                    summary_metric_col = None

                # Tabs (full center width)
                tab_labels = [qc for qc in qlist]
                first_tab_name = "Summary table" if summary_metric_col else "Results"
                tabs = st.tabs(([first_tab_name] if summary_metric_col else []) + tab_labels)

                # Summary tab
                pivot = None  # for Excel export
                if summary_metric_col:
                    long_rows = []
                    for qcode in tab_labels:
                        t = per_q_disp[qcode].copy()
                        t["Year"] = pd.to_numeric(t["Year"], errors="coerce").astype("Int64")
                        if "Demographic" not in t.columns:
                            t["Demographic"] = None
                        if summary_metric_col == "Answer1Value":
                            ans1_lab = per_q_ans1_label[qcode] or "Answer 1"
                            if ans1_lab in t.columns:
                                t["Answer1Value"] = pd.to_numeric(t[ans1_lab], errors="coerce")
                            elif "Answer 1" in t.columns:
                                t["Answer1Value"] = pd.to_numeric(t["Answer 1"], errors="coerce")
                            else:
                                t["Answer1Value"] = pd.NA
                        qlabel = f"{qcode} — {code_to_text.get(qcode, '')}".strip().rstrip(" —")
                        t["QuestionLabel"] = qlabel
                        keep = ["QuestionLabel", "Demographic", "Year", summary_metric_col]
                        long_rows.append(t[keep])
                    long_df = pd.concat(long_rows, ignore_index=True)
                    if (demo_selection != "All respondents") and (sub_selection is None) and long_df["Demographic"].notna().any():
                        idx_cols = ["QuestionLabel","Demographic"]
                    else:
                        idx_cols = ["QuestionLabel"]
                    pivot = long_df.pivot_table(index=idx_cols, columns="Year", values=summary_metric_col, aggfunc="mean")
                    pivot = pivot.reindex(selected_years, axis=1)

                    with tabs[0]:
                        header_note = summary_metric_label if summary_metric_col != "Answer1Value" else "% Answer 1 (uses each question’s Answer 1 label)"
                        st.markdown(f"### Summary table — {header_note}")
                        st.dataframe(pivot.round(1).reset_index(), use_container_width=True)
                        st.caption(
                            "Source: 2024 Public Service Employee Survey Results – Open Government Portal "
                            "https://open.canada.ca/data/en/dataset/7f625e97-9d02-4c12-a756-1ddebb50e69f"
                        )
                else:
                    st.info("Summary table unavailable because the selected questions use different answer scales/metrics. Please view detailed tabs below.")

                # Per-question tabs (details only)
                start_idx = 1 if summary_metric_col else 0
                for idx, qcode in enumerate(tab_labels, start=start_idx):
                    with tabs[idx]:
                        qtext = code_to_text.get(qcode, "")
                        st.subheader(f"{qcode} — {qtext}")
                        st.dataframe(per_q_disp[qcode], use_container_width=True)
                        st.caption(
                            "Source: 2024 Public Service Employee Survey Results – Open Government Portal "
                            "https://open.canada.ca/data/en/dataset/7f625e97-9d02-4c12-a756-1ddebb50e69f"
                        )

                # -----------------------------------
                # AI analyses (AFTER tabs)
                # -----------------------------------
                st.markdown("---")
                st.markdown("## AI Analysis")
                ai_narratives: Dict[str, str] = {}
                overall_narrative: Optional[str] = None

                if not st.session_state.get("menu1_ai_toggle", True):
                    st.info("No AI summary generated.")
                else:
                    # Per-question analyses
                    for qcode in tab_labels:
                        qtext = code_to_text.get(qcode, "")
                        metric_col = per_q_best_col[qcode]
                        metric_label = per_q_best_label[qcode]
                        with st.spinner(f"Generating Summary Analysis for {qcode}…"):
                            content, hint = _call_openai_json(
                                system=AI_SYSTEM_PROMPT,
                                user=_user_prompt_per_q(qcode, qtext, per_q_disp[qcode], metric_col, metric_label, (demo_selection != "All respondents")),
                                model=OPENAI_MODEL,
                                temperature=0.2
                            )
                        st.markdown(f"### Summary Analysis — {qcode}")
                        if content:
                            try:
                                j = json.loads(content)
                                if isinstance(j, dict) and j.get("narrative"):
                                    ai_narratives[qcode] = j["narrative"]
                                    st.write(j["narrative"])
                                    st.caption(f"Generated by OpenAI • model: {OPENAI_MODEL}")
                                else:
                                    st.caption("AI returned no narrative.")
                            except Exception:
                                st.caption("AI returned non-JSON content.")
                        else:
                            st.caption(f"AI unavailable ({hint}).")

                    # Overall analysis (only when multiple questions)
                    if len(tab_labels) > 1:
                        with st.spinner("Generating Overall Summary Analysis…"):
                            overall_payload = _user_prompt_overall_mixed(per_q_disp, per_q_text, per_q_best_col, per_q_best_label)
                            content, hint = _call_openai_json(
                                system=AI_SYSTEM_PROMPT,
                                user=overall_payload,
                                model=OPENAI_MODEL,
                                temperature=0.2
                            )
                        st.markdown("### Overall Summary Analysis")
                        if content:
                            try:
                                j = json.loads(content)
                                if isinstance(j, dict) and j.get("narrative"):
                                    overall_narrative = j["narrative"]
                                    st.write(j["narrative"])
                                    st.caption(f"Generated by OpenAI • model: {OPENAI_MODEL}")
                                else:
                                    st.caption("AI returned no narrative.")
                            except Exception:
                                st.caption("AI returned non-JSON content.")
                        else:
                            st.caption(f"AI unavailable ({hint}).")

                # -----------------------------------
                # Downloads + New search (BOTTOM)
                # -----------------------------------
                st.markdown("---")
                dl_col1, dl_col2, dl_col3 = st.columns([1, 1, 1])

                # Build Excel in-memory
                excel_bytes = None
                with io.BytesIO() as buf:
                    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
                        try:
                            if pivot is not None:
                                pivot.round(1).reset_index().to_excel(writer, sheet_name="Summary_Table", index=False)
                        except Exception:
                            pass
                        for q, df_disp in per_q_disp.items():
                            safe = q[:28]
                            df_disp.to_excel(writer, sheet_name=f"{safe}", index=False)
                        ctx = {
                            "Questions": ", ".join(question_codes),
                            "Years": ", ".join(map(str, selected_years)),
                            "Category": demo_selection,
                            "Subgroup": sub_selection or "(all in category)" if demo_selection != "All respondents" else "All respondents",
                            "Generated at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        }
                        pd.DataFrame(list(ctx.items()), columns=["Field","Value"]).to_excel(writer, sheet_name="Context", index=False)
                    excel_bytes = buf.getvalue()

                with dl_col1:
                    st.download_button(
                        label="Download data tabulations (Excel)",
                        data=excel_bytes,
                        file_name=f"PSES_multiQ_{'-'.join(map(str, selected_years))}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )

                # Build Word (.docx) with AI analyses
                def _build_ai_docx(narratives: Dict[str, str], overall_text: Optional[str]) -> Optional[bytes]:
                    try:
                        from docx import Document  # python-docx
                        from docx.shared import Pt
                    except Exception:
                        return None
                    doc = Document()
                    doc.add_heading("PSES Explorer — AI Analyses", level=1)
                    meta = doc.add_paragraph()
                    meta.add_run("Generated at: ").bold = True
                    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    meta2 = doc.add_paragraph()
                    meta2.add_run("Selection: ").bold = True
                    meta2.add_run(f"Questions: {', '.join(question_codes)}; Years: {', '.join(map(str, selected_years))}; "
                                  f"Demographic: {demo_selection}" + (f" — {sub_selection}" if sub_selection else ""))

                    # Per-question sections (in selected order)
                    for qcode in tab_labels:
                        if qcode in narratives:
                            doc.add_heading(f"Summary Analysis — {qcode}", level=2)
                            doc.add_paragraph(narratives[qcode])

                    # Overall
                    if overall_text:
                        doc.add_heading("Overall Summary Analysis", level=2)
                        doc.add_paragraph(overall_text)

                    # Footer
                    doc.add_paragraph("")
                    foot = doc.add_paragraph()
                    foot.add_run("Model: ").bold = True
                    foot.add_run(OPENAI_MODEL)
                    foot.add_run(" • Source: 2024 Public Service Employee Survey (Open Government Portal)")

                    b = io.BytesIO()
                    doc.save(b)
                    return b.getvalue()

                ai_doc_bytes = _build_ai_docx(ai_narratives, overall_narrative)

                with dl_col2:
                    if st.session_state.get("menu1_ai_toggle", True) and ai_doc_bytes:
                        st.download_button(
                            label="Download AI analysis (Word)",
                            data=ai_doc_bytes,
                            file_name=f"PSES_AI_analyses_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    else:
                        st.caption("AI analysis Word export unavailable (AI off or python-docx missing).")

                with dl_col3:
                    if st.button("Start a new search"):
                        _reset_menu1_state()
                        st.experimental_rerun()


if __name__ == "__main__":
    run_menu1()
